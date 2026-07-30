from pathlib import Path

from sqlalchemy import select
from sqlalchemy.orm import Session, selectinload

from app.core.config import settings
from app.models.review import ExportRecord, Issue, ReviewCase


class ExportService:
    def __init__(self, session: Session, output_root: Path | None = None) -> None:
        self.session = session
        self.output_root = output_root or settings.storage_root / "exports"

    def export_markdown(
        self,
        case_id: int,
        include_ai_summary: bool,
        scope: str = "final",
    ) -> Path:
        review_case = self.session.get(ReviewCase, case_id)
        if review_case is None:
            raise ValueError("Review case not found")

        issues = list(
            self.session.scalars(
                select(Issue)
                .where(Issue.case_id == case_id)
                .options(selectinload(Issue.evidence_refs))
                .order_by(Issue.risk_level.asc(), Issue.id.asc())
            ).all()
        )
        issues = self._filter_issues(issues, scope)

        target_dir = self.output_root / "cases" / str(case_id)
        target_dir.mkdir(parents=True, exist_ok=True)
        path = target_dir / f"review-case-{case_id}-v{review_case.current_version}.md"
        path.write_text(
            self._render_markdown(review_case, list(issues), include_ai_summary),
            encoding="utf-8",
        )

        self.session.add(
            ExportRecord(
                case_id=case_id,
                export_format="markdown",
                file_path=str(path),
                export_scope=scope,
            )
        )
        self.session.commit()
        return path

    def export_report(
        self,
        case_id: int,
        export_format: str,
        include_ai_summary: bool,
        scope: str = "final",
    ) -> Path:
        if export_format == "markdown":
            return self.export_markdown(case_id, include_ai_summary, scope)
        if export_format == "docx":
            return self._export_docx(case_id, include_ai_summary, scope)
        if export_format == "pdf":
            return self._export_html_pdf(case_id, include_ai_summary, scope)
        raise ValueError("Unsupported export format")

    def _export_docx(self, case_id: int, include_ai_summary: bool, scope: str) -> Path:
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except Exception:
            return self.export_markdown(case_id, include_ai_summary, scope)

        review_case = self.session.get(ReviewCase, case_id)
        if review_case is None:
            raise ValueError("Review case not found")

        issues = list(
            self.session.scalars(
                select(Issue)
                .where(Issue.case_id == case_id)
                .options(selectinload(Issue.evidence_refs))
                .order_by(Issue.risk_level.asc(), Issue.id.asc())
            ).all()
        )
        issues = self._filter_issues(issues, scope)

        document = Document()

        # Title
        title = document.add_heading(f"合同审核报告 - {review_case.title}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Meta info
        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f"版本 V{review_case.current_version} | 状态：{review_case.status}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        document.add_paragraph("")

        # Summary table
        risk_counts = {}
        for issue in issues:
            risk_counts[issue.risk_level] = risk_counts.get(issue.risk_level, 0) + 1

        summary_heading = document.add_heading("审核摘要", level=1)
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Shading Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "风险等级"
        hdr[1].text = "数量"
        risk_labels = {"high": "高风险", "medium": "中风险", "low": "低风险", "info": "提示"}
        for level in ["high", "medium", "low", "info"]:
            count = risk_counts.get(level, 0)
            if count > 0:
                row = table.add_row().cells
                row[0].text = risk_labels.get(level, level)
                row[1].text = str(count)

        document.add_paragraph("")

        # Issues
        document.add_heading("问题列表", level=1)
        risk_colors = {
            "high": RGBColor(220, 53, 69),
            "medium": RGBColor(255, 152, 0),
            "low": RGBColor(33, 150, 243),
            "info": RGBColor(100, 100, 100),
        }

        for i, issue in enumerate(issues, 1):
            # Issue heading
            h = document.add_heading(level=2)
            run = h.add_run(f"{i}. [{risk_labels.get(issue.risk_level, issue.risk_level)}] {issue.title}")
            run.font.color.rgb = risk_colors.get(issue.risk_level, RGBColor(0, 0, 0))

            # Status badge
            status_para = document.add_paragraph()
            status_run = status_para.add_run(f"状态：{issue.status} | 来源：{issue.source}")
            status_run.font.size = Pt(9)
            status_run.font.color.rgb = RGBColor(128, 128, 128)

            # Description
            document.add_paragraph(issue.description)

            # Suggestion
            if issue.suggestion:
                sug_heading = document.add_paragraph()
                sug_run = sug_heading.add_run("修改建议：")
                sug_run.bold = True
                document.add_paragraph(issue.suggestion)

            # Replacement clause
            if issue.replacement_clause:
                clause_heading = document.add_paragraph()
                clause_run = clause_heading.add_run("替代条款：")
                clause_run.bold = True
                p = document.add_paragraph(issue.replacement_clause)
                p.paragraph_format.left_indent = Inches(0.5)

            # Evidence
            if issue.evidence_refs:
                ev_heading = document.add_paragraph()
                ev_run = ev_heading.add_run("证据原文：")
                ev_run.bold = True
                for ref in issue.evidence_refs:
                    ev_text = ref.original_text or "（无原文）"
                    page_info = f"（第 {ref.page_number} 页）" if ref.page_number else ""
                    p = document.add_paragraph(f"• {ev_text} {page_info}")
                    p.paragraph_format.left_indent = Inches(0.5)

            document.add_paragraph("")

        # Footer
        document.add_paragraph("")
        footer = document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("本报告由合同审核工作台自动生成，仅供参考，请以法务最终意见为准。")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(160, 160, 160)

        target_dir = self.output_root / "cases" / str(case_id)
        target_dir.mkdir(parents=True, exist_ok=True)
        path = target_dir / f"review-case-{case_id}-v{review_case.current_version}.docx"
        document.save(str(path))

        self.session.add(
            ExportRecord(
                case_id=case_id,
                export_format="docx",
                file_path=str(path),
                export_scope=scope,
            )
        )
        self.session.commit()
        return path

    def _export_html_pdf(self, case_id: int, include_ai_summary: bool, scope: str) -> Path:
        review_case = self.session.get(ReviewCase, case_id)
        if review_case is None:
            raise ValueError("Review case not found")

        issues = list(
            self.session.scalars(
                select(Issue)
                .where(Issue.case_id == case_id)
                .options(selectinload(Issue.evidence_refs))
                .order_by(Issue.risk_level.asc(), Issue.id.asc())
            ).all()
        )
        issues = self._filter_issues(issues, scope)

        risk_labels = {"high": "高风险", "medium": "中风险", "low": "低风险", "info": "提示"}
        risk_colors = {"high": "#dc3545", "medium": "#ff9800", "low": "#2196f3", "info": "#6c757d"}

        issue_rows = ""
        for i, issue in enumerate(issues, 1):
            evidence_html = ""
            if issue.evidence_refs:
                for ref in issue.evidence_refs:
                    ev = ref.original_text or ""
                    page = f"（第 {ref.page_number} 页）" if ref.page_number else ""
                    evidence_html += f'<div class="evidence">{ev} {page}</div>'

            replacement = ""
            if issue.replacement_clause:
                replacement = f'<div class="replacement"><strong>替代条款：</strong><pre>{issue.replacement_clause}</pre></div>'

            suggestion = ""
            if issue.suggestion:
                suggestion = f'<div class="suggestion"><strong>修改建议：</strong>{issue.suggestion}</div>'

            issue_rows += f"""
            <div class="issue" style="border-left: 4px solid {risk_colors.get(issue.risk_level, '#ccc')};">
                <div class="issue-header">
                    <span class="badge" style="background:{risk_colors.get(issue.risk_level, '#ccc')};">{risk_labels.get(issue.risk_level, issue.risk_level)}</span>
                    <strong>{i}. {issue.title}</strong>
                    <span class="meta">状态：{issue.status} | 来源：{issue.source}</span>
                </div>
                <p>{issue.description}</p>
                {suggestion}
                {replacement}
                {evidence_html}
            </div>
            """

        risk_counts = {}
        for issue in issues:
            risk_counts[issue.risk_level] = risk_counts.get(issue.risk_level, 0) + 1

        summary_items = ""
        for level in ["high", "medium", "low", "info"]:
            count = risk_counts.get(level, 0)
            if count > 0:
                summary_items += f'<span class="summary-item" style="color:{risk_colors[level]};">{risk_labels[level]}: {count}</span> '

        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>合同审核报告 - {review_case.title}</title>
<style>
@page {{ margin: 2cm; size: A4; }}
body {{ font-family: "Microsoft YaHei", "PingFang SC", sans-serif; font-size: 13px; line-height: 1.6; color: #333; max-width: 800px; margin: 0 auto; padding: 20px; }}
h1 {{ text-align: center; font-size: 22px; border-bottom: 2px solid #333; padding-bottom: 10px; }}
.meta-info {{ text-align: center; color: #888; font-size: 12px; margin-bottom: 20px; }}
.summary {{ background: #f8f9fa; padding: 15px; border-radius: 6px; margin-bottom: 20px; }}
.summary-item {{ font-weight: bold; margin-right: 15px; }}
.issue {{ margin-bottom: 20px; padding: 12px; background: #fff; border-radius: 4px; }}
.issue-header {{ margin-bottom: 8px; display: flex; align-items: center; gap: 8px; }}
.badge {{ color: #fff; padding: 2px 8px; border-radius: 3px; font-size: 11px; }}
.meta {{ color: #999; font-size: 11px; margin-left: auto; }}
.evidence {{ background: #f0f0f0; padding: 6px 10px; margin: 4px 0; border-radius: 3px; font-size: 12px; }}
.suggestion {{ background: #e8f5e9; padding: 8px; border-radius: 4px; margin: 6px 0; }}
.replacement {{ background: #fff3e0; padding: 8px; border-radius: 4px; margin: 6px 0; }}
.replacement pre {{ white-space: pre-wrap; margin: 4px 0 0; }}
.footer {{ text-align: center; color: #aaa; font-size: 10px; margin-top: 40px; border-top: 1px solid #eee; padding-top: 10px; }}
</style>
</head>
<body>
<h1>合同审核报告 - {review_case.title}</h1>
<div class="meta-info">版本 V{review_case.current_version} | 状态：{review_case.status}</div>
<div class="summary">
    <strong>审核摘要：</strong> {summary_items}
</div>
<h2>问题列表</h2>
{issue_rows}
<div class="footer">本报告由合同审核工作台自动生成，仅供参考，请以法务最终意见为准。</div>
</body>
</html>"""

        target_dir = self.output_root / "cases" / str(case_id)
        target_dir.mkdir(parents=True, exist_ok=True)
        path = target_dir / f"review-case-{case_id}-v{review_case.current_version}.html"
        path.write_text(html, encoding="utf-8")

        self.session.add(
            ExportRecord(
                case_id=case_id,
                export_format="pdf",
                file_path=str(path),
                export_scope=scope,
            )
        )
        self.session.commit()
        return path

    def _filter_issues(self, issues: list[Issue], scope: str) -> list[Issue]:
        if scope == "all":
            return issues
        if scope == "high":
            return [i for i in issues if i.risk_level == "high"]
        if scope == "high_and_medium":
            return [i for i in issues if i.risk_level in {"high", "medium"}]
        if scope == "confirmed":
            return [i for i in issues if i.status == "confirmed"]
        return issues

    def _render_markdown(self, review_case, issues, include_ai_summary) -> str:
        risk_labels = {"high": "高风险", "medium": "中风险", "low": "低风险", "info": "提示"}
        lines = [
            f"# 合同审核报告 - {review_case.title}",
            "",
            f"版本：V{review_case.current_version}",
            f"状态：{review_case.status}",
            "",
        ]

        risk_counts = {}
        for issue in issues:
            risk_counts[issue.risk_level] = risk_counts.get(issue.risk_level, 0) + 1
        lines.append("## 审核摘要")
        for level in ["high", "medium", "low", "info"]:
            count = risk_counts.get(level, 0)
            if count:
                lines.append(f"- {risk_labels[level]}: {count} 个")
        lines.append("")

        lines.append("## 问题列表")
        for i, issue in enumerate(issues, 1):
            lines.append(f"### {i}. [{risk_labels.get(issue.risk_level, issue.risk_level)}] {issue.title}")
            lines.append(f"状态：{issue.status} | 来源：{issue.source}")
            lines.append("")
            lines.append(issue.description)
            if issue.suggestion:
                lines.append(f"\n**修改建议：** {issue.suggestion}")
            if issue.replacement_clause:
                lines.append(f"\n**替代条款：**\n```\n{issue.replacement_clause}\n```")
            if issue.evidence_refs:
                lines.append("\n**证据原文：**")
                for ref in issue.evidence_refs:
                    page = f"（第 {ref.page_number} 页）" if ref.page_number else ""
                    lines.append(f"- {ref.original_text or '（无原文）'} {page}")
            lines.append("")

        lines.append("")
        lines.append("---")
        lines.append("*本报告由合同审核工作台自动生成，仅供参考，请以法务最终意见为准。*")
        return "\n".join(lines)
